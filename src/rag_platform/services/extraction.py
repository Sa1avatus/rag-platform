import email
import io
import json
import re
from dataclasses import dataclass
from email import policy
from html.parser import HTMLParser
from pathlib import PurePosixPath

from rag_platform.services.source_safety import ArchiveMember, inspect_zip, validate_document


@dataclass(frozen=True)
class ExtractedDocument:
    filename: str
    content: str


class _TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self._ignored_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style"}:
            self._ignored_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style"} and self._ignored_depth:
            self._ignored_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self._ignored_depth:
            self.parts.append(data)


def extract(filename: str, content: bytes, mime_type: str | None = None) -> list[ExtractedDocument]:
    suffix = validate_document(filename, content, mime_type)
    if suffix == ".zip":
        return [_extract_member(member) for member in inspect_zip(content)]
    return [ExtractedDocument(filename, _extract_content(suffix, content))]


def _extract_member(member: ArchiveMember) -> ExtractedDocument:
    suffix = PurePosixPath(member.path).suffix.lower()
    return ExtractedDocument(member.path, _extract_content(suffix, member.content))


def _extract_content(suffix: str, content: bytes) -> str:
    if suffix == ".pdf":
        return _extract_pdf(content)
    if suffix == ".docx":
        return _extract_docx(content)
    if suffix == ".eml":
        return _extract_email(content)
    decoded = content.decode("utf-8-sig", errors="replace")
    if suffix == ".json":
        value = json.loads(decoded)
        return json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True)
    if suffix in {".html", ".htm"}:
        parser = _TextExtractor()
        parser.feed(decoded)
        return _normalize(" ".join(parser.parts))
    return _normalize(decoded)


def _extract_email(content: bytes) -> str:
    message = email.message_from_bytes(content, policy=policy.default)
    fields = [
        f"Subject: {message.get('subject', '')}",
        f"From: {message.get('from', '')}",
        f"To: {message.get('to', '')}",
    ]
    bodies: list[str] = []
    for part in message.walk():
        is_plain_body = (
            part.get_content_type() == "text/plain"
            and part.get_content_disposition() != "attachment"
        )
        if is_plain_body:
            bodies.append(part.get_content())
    return _normalize("\n".join([*fields, *bodies]))


def _extract_pdf(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    return _normalize("\n\n".join(page.extract_text() or "" for page in reader.pages))


def _extract_docx(content: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(content))
    return _normalize("\n".join(paragraph.text for paragraph in document.paragraphs))


def _normalize(value: str) -> str:
    value = value.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    value = re.sub(r"[\t ]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()
